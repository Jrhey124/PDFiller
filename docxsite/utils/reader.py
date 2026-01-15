from datetime import date
import json
import os

from docx import Document
from docx.text.run import Run
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
json_path = os.path.join(BASE_DIR, 'resources.json')

with open(json_path, 'r', encoding='utf-8') as f:
    resources = json.load(f)


# ---------- Helpers ----------

def is_image_run(run: Run) -> bool:
    """Return True if this run contains an inline image (<w:drawing>)."""
    return bool(run._r.findall(qn('w:drawing')))


def copy_run_style(src: Run, dest: Run):
    """Copy bold, italic, underline, font name/size/color from src→dest."""
    dest.bold      = src.bold
    dest.italic    = src.italic
    dest.underline = src.underline
    dest.font.name = src.font.name
    dest.font.size = src.font.size
    if src.font.color.rgb:
        dest.font.color.rgb = src.font.color.rgb


def insert_run_after(paragraph, anchor_run: Run, text: str, style_src: Run) -> Run:
    """
    Insert a new run with `text` immediately after `anchor_run`,
    then copy styling from `style_src`.
    """
    new_run = paragraph.add_run(text)
    anchor_elm = anchor_run._r
    new_elm    = new_run._r
    anchor_elm.addnext(new_elm)
    copy_run_style(style_src, new_run)
    return new_run


def merge_adjacent_runs(paragraph):
    """Collapse adjacent runs with identical style, skipping image runs."""
    def same_style(a, b):
        return (
            a.bold      == b.bold
        and a.italic    == b.italic
        and a.underline == b.underline
        and a.font.name == b.font.name
        and a.font.size == b.font.size
        and (getattr(a.font.color, "rgb", None)
             == getattr(b.font.color, "rgb", None))
        )
    i = 0
    while i < len(paragraph.runs) - 1:
        run, nxt = paragraph.runs[i], paragraph.runs[i+1]
        if is_image_run(run) or is_image_run(nxt):
            i += 1
            continue
        if same_style(run, nxt):
            run.text += nxt.text
            nxt._r.getparent().remove(nxt._r)
        else:
            i += 1


def remove_tab_elements(run):
    """Remove all <w:tab/> elements from this run."""
    for tab in run._r.findall(qn('w:tab')):
        run._r.remove(tab)


def replace_tag_in_paragraph(paragraph, tag: str, replacement: str, preserve_tabs: bool = False):
    """Replace placeholders in text runs only, leaving image runs untouched."""
    merge_adjacent_runs(paragraph)
    while True:
        for run in paragraph.runs:
            if is_image_run(run):
                continue
            if tag in run.text:
                if preserve_tabs:
                    for text_elem in run._r.findall(qn('w:t')):
                        if text_elem.text and tag in text_elem.text:
                            text_elem.text = text_elem.text.replace(tag, replacement)
                else:
                    remove_tab_elements(run)
                    before, _, after = run.text.partition(tag)
                    run.text = before
                    repl = insert_run_after(paragraph, run, replacement, run)
                    insert_run_after(paragraph, repl, after, run)
                break
        else:
            return


def smart_replace_in_docx(input_path: str, replacements: dict, preserve_tabs_for: list = None):
    """
    Open a .docx at input_path, replace all tags in `replacements`,
    preserving styling, and return the Document object.
    """
    doc = Document(input_path)
    if preserve_tabs_for is None:
        preserve_tabs_for = []

    # Body paragraphs
    for p in doc.paragraphs:
        for tag, val in replacements.items():
            replace_tag_in_paragraph(p, tag, val, preserve_tabs=(tag in preserve_tabs_for))

    # Tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for tag, val in replacements.items():
                        replace_tag_in_paragraph(p, tag, val, preserve_tabs=(tag in preserve_tabs_for))

    # Headers & Footers
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            for p in hdr.paragraphs:
                for tag, val in replacements.items():
                    replace_tag_in_paragraph(p, tag, val, preserve_tabs=(tag in preserve_tabs_for))

    return doc


# ---------- Main Request Letter ----------

def Request_Letter(addressee, adviser, students, course_section, details, gender):
    """Build a .docx letter by replacing placeholders, preserving style and images."""
    gender = "M" if gender == "Mr." else "F" if gender == "Ms." else "N"

    addressees_list = resources.get("addressees", {})
    if addressee not in addressees_list:
        raise ValueError(f"Addressee '{addressee}' not found in resources.json")

    replacements = {
        "<DATE>": date.today().strftime("%B %d, %Y"),
        "<TITLE AND NAME OF ADDRESSEE>": addressee,
        "<DESIGNATION>": addressees_list[addressee].get("designation", ""),
        "<NAME OF OFFICE>": addressees_list[addressee].get("office_name", ""),
        "<ADDRESS OF OFFICE>": addressees_list[addressee].get("office_address", ""),
        "<STUDENTS>": (students[0] if len(students) == 1 else ", ".join(students)),
        "<IS_ARE>": "is" if len(students) == 1 else "are",
        "<POSSESSIVE_PRONOUN>": (
            "their" if len(students) > 1
            else "her" if gender == "F"
            else "his" if gender == "M"
            else "their"
        ),
        "<COURSE_SECTION>": course_section,
        "<OBJECT_PRONOUN>": (
            "them" if len(students) > 1
            else "her" if gender == "F"
            else "him" if gender == "M"
            else "them"
        ),
        "<DATA_GATHERING_DETAILS>": "" if not details else f"({details})",
        "<SUBJECT_PRONOUN>": (
            "they" if len(students) > 1
            else "she" if gender == "F"
            else "he" if gender == "M"
            else "they"
        ),
        "<ADVISER>": adviser,
        "<STUDENT_FULLNAME>": (
            students[0] if len(students) == 1
            else students[0] + " et al."
        )
    }

    template_path = os.path.join(BASE_DIR, 'Templates', 'ORIGINAL FORMAT of Request Letter (026).docx')
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")

    doc = smart_replace_in_docx(template_path, replacements, preserve_tabs_for=['<STUDENT_FULLNAME>'])
    return doc


# ---------- Test harness ----------

if __name__ == "__main__":
    Request_Letter(
        addressee="Dr. Maria C. Santos",
        adviser="Engr. Berlim Limbauan",
        students=["Juan Dela Cruz", "Pedro Penduko", "Maria Clara"],
        course_section="BSCE 3-1",
        details="for our research project on sustainable building",
        gender="N"
    )
